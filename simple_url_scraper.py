#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Simple WeChat URL Scraper
简单微信文章URL处理工具

直接处理用户提供的微信文章URL，获取内容并保存为飞书格式
"""

import time
import random
import json
import re
import base64
import requests
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

from typing import List, Dict, Any, Optional
from loguru import logger
from datetime import datetime
import os

from utils import clean_text




class SimpleUrlScraper:
    """简单URL处理工具"""
    
    def __init__(self, headless: bool = False):
        """
        初始化工具 - 快速模式，不立即启动浏览器
        
        Args:
            headless: 是否无头模式
        """
        self.driver = None
        self.headless = headless
        
        logger.info("简单URL处理工具初始化完成（浏览器将懒加载）")
    
    def setup_browser(self, headless: bool = True) -> Optional[webdriver.Chrome]:
        """设置Chrome浏览器 - 针对exe环境优化"""
        try:
            logger.info("🚀 正在初始化浏览器（exe优化模式）...")
            
            # Chrome选项配置 - 针对exe环境优化
            chrome_options = Options()
            
            # 基础选项
            if headless:
                chrome_options.add_argument("--headless=new")  # 使用新的headless模式
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--disable-extensions")
            chrome_options.add_argument("--window-size=1920,1080")
            
            # 重要：不禁用JavaScript和CSS，验证码需要这些
            # chrome_options.add_argument("--disable-javascript")  # 注释掉
            # chrome_options.add_argument("--disable-css")  # 注释掉
            
            # 反检测配置 - 针对exe环境
            chrome_options.add_argument("--disable-blink-features=AutomationControlled")
            chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
            chrome_options.add_experimental_option('useAutomationExtension', False)
            
            # 模拟真实浏览器环境
            chrome_options.add_argument("--disable-web-security")
            chrome_options.add_argument("--allow-running-insecure-content")
            chrome_options.add_argument("--disable-features=VizDisplayCompositor")
            
            # exe环境特殊配置
            import sys
            if getattr(sys, 'frozen', False):  # 检测是否为exe环境
                logger.info("🔧 检测到exe环境，应用特殊配置...")
                chrome_options.add_argument("--disable-logging")
                chrome_options.add_argument("--disable-gpu-logging")
                chrome_options.add_argument("--silent")
                chrome_options.add_argument("--no-first-run")
                chrome_options.add_argument("--no-default-browser-check")
            
            # GitHub Actions特殊配置
            if os.getenv('GITHUB_ACTIONS'):
                logger.info("🔧 检测到GitHub Actions环境，应用特殊配置...")
                chrome_options.add_argument("--no-first-run")
                chrome_options.add_argument("--no-default-browser-check")
                chrome_options.add_argument("--disable-default-apps")
                chrome_options.add_argument("--disable-background-timer-throttling")
                chrome_options.add_argument("--disable-backgrounding-occluded-windows")
                chrome_options.add_argument("--disable-renderer-backgrounding")
                chrome_options.add_argument("--disable-features=TranslateUI")
                chrome_options.add_argument("--disable-ipc-flooding-protection")
                chrome_options.add_argument("--single-process")
                chrome_options.add_argument("--remote-debugging-port=9222")
                
                # 设置用户数据目录
                user_data_dir = "/tmp/chrome-user-data"
                os.makedirs(user_data_dir, exist_ok=True)
                chrome_options.add_argument(f"--user-data-dir={user_data_dir}")
            
            # 用户代理 - 使用Windows Chrome
            user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            chrome_options.add_argument(f"--user-agent={user_agent}")
            
            # 页面加载策略
            chrome_options.page_load_strategy = 'normal'  # 改为normal以确保验证码加载
            
            # 超时设置
            chrome_options.add_argument("--timeout=60000")  # 延长超时时间
            
            # 日志配置
            chrome_options.add_argument("--log-level=3")
            
            # 性能优化 - 但不禁用图片（验证码需要）
            prefs = {
                "profile.default_content_setting_values": {
                    "plugins": 2,
                    "popups": 2,
                    "geolocation": 2,
                    "notifications": 2,
                    "media_stream": 2,
                },
                # 不禁用图片，验证码需要图片加载
                "profile.managed_default_content_settings": {
                    # "images": 2  # 注释掉，允许图片加载
                }
            }
            chrome_options.add_experimental_option("prefs", prefs)
            
            # WebDriver Manager配置
            service = None
            try:
                logger.info("⚡ 启动Chrome浏览器...")
                
                # 在GitHub Actions环境中，尝试使用系统Chrome
                if os.getenv('GITHUB_ACTIONS'):
                    # 检查系统Chrome路径
                    chrome_paths = [
                        "/usr/bin/google-chrome",
                        "/usr/bin/google-chrome-stable",
                        "/usr/bin/chromium-browser",
                        "/snap/bin/chromium"
                    ]
                    
                    chrome_binary = None
                    for path in chrome_paths:
                        if os.path.exists(path):
                            chrome_binary = path
                            logger.info(f"🔍 找到Chrome二进制文件: {chrome_binary}")
                            break
                    
                    if chrome_binary:
                        chrome_options.binary_location = chrome_binary
                    
                    # 尝试使用系统chromedriver
                    chromedriver_paths = [
                        "/usr/bin/chromedriver",
                        "/usr/local/bin/chromedriver"
                    ]
                    
                    chromedriver_path = None
                    for path in chromedriver_paths:
                        if os.path.exists(path):
                            chromedriver_path = path
                            logger.info(f"🔍 找到ChromeDriver: {chromedriver_path}")
                            break
                    
                    if chromedriver_path:
                        service = Service(chromedriver_path)
                    else:
                        # 使用WebDriverManager作为后备
                        from webdriver_manager.chrome import ChromeDriverManager
                        service = Service(ChromeDriverManager().install())
                else:
                    # 本地环境使用WebDriverManager
                    from webdriver_manager.chrome import ChromeDriverManager
                    service = Service(ChromeDriverManager().install())
                
                # 创建WebDriver实例
                if service:
                    driver = webdriver.Chrome(service=service, options=chrome_options)
                else:
                    driver = webdriver.Chrome(options=chrome_options)
                
                # 设置超时
                driver.set_page_load_timeout(60)  # 延长页面加载超时
                driver.implicitly_wait(15)  # 延长隐式等待
                
                # 反检测脚本注入
                driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
                driver.execute_script("Object.defineProperty(navigator, 'plugins', {get: () => [1, 2, 3, 4, 5]})")
                driver.execute_script("Object.defineProperty(navigator, 'languages', {get: () => ['zh-CN', 'zh', 'en']})")
                
                logger.info("✅ Chrome浏览器启动成功（已优化验证码支持）")
                return driver
                
            except Exception as e:
                logger.error(f"浏览器初始化失败: {e}")
                
                # 尝试降级方案
                if not os.getenv('GITHUB_ACTIONS'):
                    logger.info("🔄 尝试降级方案...")
                    try:
                        # 移除一些可能有问题的选项
                        chrome_options = Options()
                        chrome_options.add_argument("--headless")
                        chrome_options.add_argument("--no-sandbox")
                        chrome_options.add_argument("--disable-dev-shm-usage")
                        
                        driver = webdriver.Chrome(options=chrome_options)
                        driver.set_page_load_timeout(30)
                        logger.info("✅ 降级方案成功")
                        return driver
                    except Exception as e2:
                        logger.error(f"降级方案也失败: {e2}")
                
                return None
                
        except Exception as e:
            logger.error(f"设置浏览器时出错: {e}")
            return None
    
    def extract_article_info(self, url: str) -> dict:
        """
        快速提取文章信息 - 优化版本
        """
        try:
            if not self.driver:
                self.driver = self.setup_browser(self.headless)
                if not self.driver:
                    return {"error": "浏览器初始化失败"}
            
            logger.info(f"正在访问URL: {url}")
            
            # 快速访问页面
            self.driver.get(url)
            
            # 快速等待（最多2秒）
            try:
                wait = WebDriverWait(self.driver, 2)
                wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            except TimeoutException:
                logger.warning("页面加载超时，尝试继续提取")
            
            # 快速提取文章信息
            article_info = self._extract_article_content()
            
            if article_info and 'error' not in article_info:
                return article_info
            else:
                return {"error": "文章信息提取失败"}
                
        except Exception as e:
            logger.error(f"提取文章信息失败: {e}")
            return {"error": f"提取失败: {str(e)}"}
    
    def save_as_pdf(self, url: str, output_path: str) -> bool:
        """
        保存URL为PDF文件（原始实现，无重试机制）
        """
        try:
            logger.info(f"正在保存PDF: {url}")
            # 1. 使用Selenium加载完整页面内容
            if not self.driver:
                self.driver = self.setup_browser(headless=True)
            self.driver.get(url)
            time.sleep(2)
            # 等待页面加载
            self._wait_for_basic_page_load()
            self._human_like_scroll_and_load()
            self.driver.execute_script("window.scrollTo(0, 0);")
            time.sleep(0.2)
            # 注入CSS样式
            css_style = """
            <style>
                @page { margin: 0 !important; padding: 0 !important; size: A4 !important; }
                * { box-sizing: border-box !important; }
                body, html { margin: 0 !important; padding: 0 !important; width: 100% !important; max-width: 100% !important; }
                #js_content, .rich_media_content, .rich_media_area_primary { margin: 0 !important; padding: 5px !important; width: 100% !important; max-width: 100% !important; }
                .rich_media_meta_list, .rich_media_tool, .qr_code_pc_outer, .reward_qrcode_area, .reward_area, #js_pc_qr_code_img, .function_mod, .profile_container, .rich_media_global_msg { display: none !important; }
            </style>
            """
            self.driver.execute_script(f"""
                var style = document.createElement('style');
                style.type = 'text/css';
                style.innerHTML = `{css_style}`;
                document.head.appendChild(style);
            """)
            time.sleep(0.3)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            pdf_options = {
                'paperFormat': 'A4',
                'printBackground': True,
                'marginTop': 0,
                'marginBottom': 0,
                'marginLeft': 0,
                'marginRight': 0,
                'preferCSSPageSize': True,
                'displayHeaderFooter': False,
                'scale': 1.0,
                'landscape': False,
                'transferMode': 'ReturnAsBase64',
                'generateTaggedPDF': False
            }
            pdf_data = self.driver.execute_cdp_cmd('Page.printToPDF', pdf_options)
            with open(output_path, 'wb') as f:
                f.write(base64.b64decode(pdf_data['data']))
            logger.success(f"PDF保存成功: {output_path}")
            return True
        except Exception as e:
            logger.error(f"保存PDF失败: {e}")
            return False
    
    def save_as_docx(self, url: str, output_path: str) -> bool:
        """
        保存URL为Word文档 - 完整内容版本
        使用Selenium动态加载 + BeautifulSoup解析确保内容完整性
        """
        try:
            logger.info(f"正在保存Word文档: {url}")
            
            # 1. 使用Selenium加载完整页面内容（包括动态内容）
            article_data = self._extract_wechat_article_with_selenium(url)
            
            if not article_data or 'error' in article_data:
                logger.error("无法获取文章内容")
                return False
            
            # 2. 创建Word文档
            doc = Document()
            
            # 设置默认字体样式（微软雅黑，10.5号字）
            try:
                from docx.enum.style import WD_STYLE_TYPE
                style = doc.styles.add_style('DefaultParagraph', WD_STYLE_TYPE.PARAGRAPH)
                font = style.font
                font.name = '微软雅黑'
                font.size = Pt(10.5)
            except:
                logger.debug("设置默认字体样式失败，使用系统默认")
            
            # 3. 添加标题
            title = article_data.get('title', '微信文章')
            title_paragraph = doc.add_heading(title, level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 4. 添加作者和发布时间
            meta_info = f"作者: {article_data.get('author', '未知')}\n"
            meta_info += f"发布时间: {article_data.get('publish_date', '未知')}\n"
            meta_info += f"原文链接: {url}\n"
            
            meta_paragraph = doc.add_paragraph(meta_info)
            meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 5. 添加分隔线
            doc.add_paragraph("=" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 6. 处理HTML内容，确保所有信息都包含
            content_soup = article_data.get('content_soup')
            images = article_data.get('images', [])
            
            if content_soup:
                logger.info(f"🔄 开始处理文档内容，包含 {len(images)} 张图片")
                try:
                    # 添加超时保护的内容处理
                    import signal
                    
                    def timeout_handler(signum, frame):
                        raise TimeoutError("文档处理超时")
                    
                    # 设置30秒超时
                    if hasattr(signal, 'SIGALRM'):  # Unix系统
                        signal.signal(signal.SIGALRM, timeout_handler)
                        signal.alarm(30)
                    
                    self._process_wechat_content_to_docx(doc, content_soup, images)
                    
                    if hasattr(signal, 'SIGALRM'):
                        signal.alarm(0)  # 取消超时
                        
                except TimeoutError:
                    logger.warning("⏰ 文档处理超时，使用简化处理")
                    # 降级到简单文本处理
                    text_content = content_soup.get_text(strip=True)
                    if text_content:
                        doc.add_paragraph(text_content[:5000])  # 限制长度
                except Exception as e:
                    logger.warning(f"⚠️ 内容处理异常: {e}，使用简化处理")
                    # 降级到简单文本处理
                    try:
                        text_content = content_soup.get_text(strip=True)
                        if text_content:
                            doc.add_paragraph(text_content[:5000])  # 限制长度
                    except:
                        doc.add_paragraph("内容提取失败")
            else:
                logger.warning("没有找到文章内容")
                doc.add_paragraph("未能提取到文章内容")
            
            # 7. 确保输出目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 8. 保存文档
            try:
                doc.save(output_path)
                logger.success(f"Word文档保存成功: {output_path}")
                logger.info(f"📊 统计信息: 图片={len(images)}张, 文件大小={os.path.getsize(output_path)/1024:.1f}KB")
                return True
            except Exception as e:
                logger.error(f"保存Word文档文件失败: {e}")
                return False
            
        except Exception as e:
            logger.error(f"保存Word文档失败: {e}")
            import traceback
            logger.debug(f"详细错误信息: {traceback.format_exc()}")
            return False
    
    def _extract_wechat_article_with_selenium(self, url: str) -> dict:
        """
        使用Selenium获取完整的微信文章内容（包括动态加载内容）
        """
        try:
            logger.info(f"🚀 使用Selenium获取完整文章内容: {url}")
            
            # 确保浏览器已初始化
            if not self.driver:
                self.driver = self.setup_browser(self.headless)
                if not self.driver:
                    return {"error": "浏览器初始化失败"}
            
            # 访问页面
            self.driver.get(url)
            
            # 等待页面基础加载
            try:
                wait = WebDriverWait(self.driver, 15)
                wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
                logger.debug("✅ 页面基础框架加载完成")
            except TimeoutException:
                logger.warning("⏰ 页面加载超时，尝试继续")
            
            # 处理验证码（如果存在）
            if not self._handle_captcha_if_present():
                logger.error("❌ 验证码处理失败")
                return {"error": "验证码处理失败"}
            
            # 等待内容加载
            self._wait_for_basic_page_load()
            
            # 完整滚动加载所有动态内容
            self._human_like_scroll_and_load()
            
            # 额外等待确保所有动态内容都已加载
            logger.info("⏳ 等待动态内容完全加载...")
            time.sleep(3)
            
            # 获取完全渲染后的页面HTML
            page_source = self.driver.page_source
            
            # 使用BeautifulSoup解析完整的HTML
            soup = BeautifulSoup(page_source, 'html.parser')
            
            # 调试：保存完整HTML
            debug_html_path = "debug_selenium_page.html"
            with open(debug_html_path, 'w', encoding='utf-8') as f:
                f.write(page_source)
            logger.debug(f"完整HTML已保存到: {debug_html_path}")
            
            # 提取标题 - 从渲染后的页面
            title = self._extract_title_from_selenium_soup(soup)
            
            # 提取作者
            author = self._extract_author_from_selenium_soup(soup)
            
            # 提取发布时间
            publish_date = self._extract_publish_date_from_selenium_soup(soup)
            
            # 提取文章正文内容 - 从完全加载的页面
            content_soup = self._extract_content_from_selenium_soup(soup)
            
            if not content_soup:
                return {"error": "未找到文章正文内容"}
            
            # 下载图片（使用Selenium已经加载的图片）
            images = self._download_images_from_selenium_soup(content_soup, url)
            
            logger.success(f"✅ 成功提取完整文章内容: {title[:30]}...")
            logger.info(f"📊 内容统计: 段落数={len(content_soup.find_all(['p', 'div']))}, 图片数量={len(images)}")
            
            return {
                'title': title,
                'author': author,
                'publish_date': publish_date,
                'content_soup': content_soup,
                'images': images,
                'url': url
            }
            
        except Exception as e:
            logger.error(f"Selenium提取文章失败: {e}")
            return {"error": f"提取失败: {str(e)}"}
    
    def _extract_title_from_selenium_soup(self, soup: BeautifulSoup) -> str:
        """从Selenium渲染的页面中提取标题"""
        try:
            # 微信文章标题的选择器（渲染后）
            title_selectors = [
                'h1#activity-name',
                'h2.rich_media_title',
                '.rich_media_title',
                'h1.rich_media_title',
                '.rich_media_title h1',
                '#activity-name',
                '.wx_title',
                '[data-role="title"]',
                'h1',
                'title'
            ]
            
            for selector in title_selectors:
                title_elem = soup.select_one(selector)
                if title_elem and title_elem.get_text(strip=True):
                    title = title_elem.get_text(strip=True)
                    if len(title) > 3:  # 确保不是空标题
                        logger.debug(f"通过选择器 {selector} 找到标题: {title[:50]}...")
                        return title
            
            # 从页面标题提取
            title_tag = soup.find('title')
            if title_tag and title_tag.get_text(strip=True):
                title = title_tag.get_text(strip=True)
                if '微信公众平台' not in title and len(title) > 3:
                    logger.debug(f"从title标签找到标题: {title[:50]}...")
                    return title
            
            # 从meta标签提取
            meta_title = soup.find('meta', property='og:title')
            if meta_title and meta_title.get('content'):
                title = meta_title.get('content').strip()
                if title and len(title) > 3:
                    logger.debug(f"从meta标签找到标题: {title[:50]}...")
                    return title
            
            return "未知标题"
            
        except Exception as e:
            logger.debug(f"提取标题失败: {e}")
            return "未知标题"
    
    def _extract_author_from_selenium_soup(self, soup: BeautifulSoup) -> str:
        """从Selenium渲染的页面中提取作者"""
        try:
            # 微信文章作者的选择器（渲染后）
            author_selectors = [
                '.rich_media_meta_nickname',
                '.profile_nickname',
                '[data-role="nickname"]',
                '.account_nickname',
                '.wx_author',
                '#profileBt .nickname'
            ]
            
            for selector in author_selectors:
                author_elem = soup.select_one(selector)
                if author_elem and author_elem.get_text(strip=True):
                    author = author_elem.get_text(strip=True)
                    logger.debug(f"通过选择器 {selector} 找到作者: {author}")
                    return author
            
            return "未知作者"
            
        except Exception as e:
            logger.debug(f"提取作者失败: {e}")
            return "未知作者"
    
    def _extract_publish_date_from_selenium_soup(self, soup: BeautifulSoup) -> str:
        """从Selenium渲染的页面中提取发布时间"""
        try:
            # 发布时间的选择器（渲染后）
            date_selectors = [
                '#publish_time',
                '.rich_media_meta_text',
                '[data-role="publish-time"]',
                '.wx_time',
                '.time'
            ]
            
            for selector in date_selectors:
                date_elem = soup.select_one(selector)
                if date_elem and date_elem.get_text(strip=True):
                    date_text = date_elem.get_text(strip=True)
                    logger.debug(f"通过选择器 {selector} 找到发布时间: {date_text}")
                    return date_text
            
            return "未知时间"
            
        except Exception as e:
            logger.debug(f"提取发布时间失败: {e}")
            return "未知时间"
    
    def _extract_content_from_selenium_soup(self, soup: BeautifulSoup) -> BeautifulSoup:
        """从Selenium渲染的页面中提取文章正文区域"""
        try:
            # 微信文章正文的选择器（渲染后）
            content_selectors = [
                '#js_content',
                '.rich_media_content',
                '.appmsg_content_text',
                '[data-role="article-content"]',
                '.rich_media_area_primary',
                '.rich_media_wrp',
                '#js_content_container'
            ]
            
            for selector in content_selectors:
                content_elem = soup.select_one(selector)
                if content_elem:
                    # 检查内容是否足够多
                    text_content = content_elem.get_text(strip=True)
                    if text_content and len(text_content) > 100:  # 至少100个字符
                        logger.debug(f"✅ 通过选择器 {selector} 找到正文内容，长度: {len(text_content)}")
                        return content_elem
            
            # 如果没有找到合适的容器，智能提取
            logger.warning("未找到专用正文容器，尝试智能提取")
            
            # 查找包含最多文本内容的div
            all_divs = soup.find_all('div')
            best_div = None
            max_text_length = 0
            
            for div in all_divs:
                text = div.get_text(strip=True)
                if len(text) > max_text_length and len(text) > 200:  # 至少200个字符
                    # 排除导航、头部、脚部等区域
                    div_class = ' '.join(div.get('class', []))
                    div_id = div.get('id', '')
                    
                    exclude_keywords = ['nav', 'header', 'footer', 'sidebar', 'menu', 'toolbar', 'ad']
                    if not any(keyword in div_class.lower() + div_id.lower() for keyword in exclude_keywords):
                        max_text_length = len(text)
                        best_div = div
            
            if best_div:
                logger.debug(f"智能提取找到内容区域，文本长度: {max_text_length}")
                return best_div
            
            logger.warning("使用整个body作为内容")
            return soup.find('body') or soup
            
        except Exception as e:
            logger.error(f"提取正文内容失败: {e}")
            return None
    
    def _download_images_from_selenium_soup(self, content_soup: BeautifulSoup, base_url: str) -> list:
        """从Selenium渲染的页面中下载图片（优化版）"""
        try:
            images_info = []
            img_tags = content_soup.find_all('img')
            
            if not img_tags:
                logger.info("📷 未发现图片")
                return images_info
            
            logger.info(f"🖼️ 发现 {len(img_tags)} 张图片，开始下载...")
            
            # 创建图片下载目录
            import urllib.parse
            parsed_url = urllib.parse.urlparse(base_url)
            safe_domain = re.sub(r'[<>:"/\\|?*]', '_', parsed_url.netloc)
            img_dir = os.path.join("output", "images", safe_domain)
            os.makedirs(img_dir, exist_ok=True)
            
            # 限制同时下载的图片数量，避免卡住
            max_images = min(len(img_tags), 20)  # 最多下载20张图片
            successful_downloads = 0
            
            for i, img in enumerate(img_tags[:max_images]):
                try:
                    # 从Selenium加载的页面中，src应该已经被完全解析
                    img_src = img.get('src') or img.get('data-src') or img.get('data-original')
                    if not img_src:
                        continue
                    
                    # 处理完整的URL
                    if img_src.startswith('//'):
                        img_src = 'https:' + img_src
                    elif img_src.startswith('/'):
                        img_src = f"https://{parsed_url.netloc}" + img_src
                    elif not img_src.startswith('http'):
                        img_src = urllib.parse.urljoin(base_url, img_src)
                    
                    # 跳过过小的图片（可能是图标）或base64图片
                    if ('icon' in img_src.lower() or 
                        'logo' in img_src.lower() or 
                        img_src.startswith('data:') or
                        len(img_src) > 1000):  # 跳过超长URL
                        continue
                    
                    # 生成本地文件名
                    img_filename = f"img_{i+1:03d}.jpg"
                    img_path = os.path.join(img_dir, img_filename)
                    
                    # 下载图片（带超时控制）
                    download_success = self._download_image_with_timeout(img_src, img_path, timeout=10)
                    
                    if download_success:
                        images_info.append({
                            'url': img_src,
                            'local_path': img_path,
                            'filename': img_filename
                        })
                        successful_downloads += 1
                        logger.debug(f"📷 下载图片成功: {img_filename}")
                    else:
                        # 即使下载失败，也记录图片信息，但local_path为None
                        images_info.append({
                            'url': img_src,
                            'local_path': None,
                            'filename': img_filename
                        })
                        logger.debug(f"📷 下载图片失败: {img_src[:50]}...")
                
                except Exception as e:
                    logger.debug(f"处理图片时出错: {e}")
                    continue
            
            logger.info(f"🖼️ 图片下载完成: {successful_downloads}/{max_images} 张成功")
            return images_info
            
        except Exception as e:
            logger.warning(f"图片下载过程异常: {e}")
            return []

    def _download_image_with_timeout(self, img_url: str, save_path: str, timeout: int = 10) -> bool:
        """带超时控制的图片下载"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Referer': 'https://mp.weixin.qq.com/',
                'Accept': 'image/webp,image/apng,image/*,*/*;q=0.8'
            }
            
            # 使用较短的超时时间
            response = requests.get(img_url, headers=headers, timeout=timeout, stream=True)
            response.raise_for_status()
            
            # 检查内容类型
            content_type = response.headers.get('content-type', '').lower()
            if not any(img_type in content_type for img_type in ['image/', 'jpeg', 'jpg', 'png', 'gif', 'webp']):
                logger.debug(f"非图片内容类型: {content_type}")
                return False
            
            # 检查文件大小，避免下载过大的文件
            content_length = response.headers.get('content-length')
            if content_length and int(content_length) > 10 * 1024 * 1024:  # 10MB限制
                logger.debug(f"图片文件过大: {content_length} bytes")
                return False
            
            # 写入文件
            with open(save_path, 'wb') as f:
                downloaded_size = 0
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
                        downloaded_size += len(chunk)
                        # 限制下载大小
                        if downloaded_size > 10 * 1024 * 1024:  # 10MB限制
                            logger.debug("下载文件过大，停止下载")
                            return False
            
            # 验证文件是否成功下载
            if os.path.exists(save_path) and os.path.getsize(save_path) > 0:
                return True
            else:
                return False
            
        except requests.exceptions.Timeout:
            logger.debug(f"图片下载超时: {img_url}")
            return False
        except requests.exceptions.RequestException as e:
            logger.debug(f"图片下载网络错误: {e}")
            return False
        except Exception as e:
            logger.debug(f"图片下载异常: {e}")
            return False
    
    def _extract_wechat_article_by_requests(self, url: str) -> dict:
        """
        使用Requests获取微信文章内容
        """
        try:
            logger.info(f"🌐 使用Requests获取文章内容: {url}")
            
            # 设置请求头绕过反爬机制
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                'Accept-Encoding': 'gzip, deflate, br',
                'DNT': '1',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
                'Sec-Fetch-Dest': 'document',
                'Sec-Fetch-Mode': 'navigate',
                'Sec-Fetch-Site': 'none',
                'Cache-Control': 'max-age=0'
            }
            
            # 发送请求
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            response.encoding = 'utf-8'
            
            # 解析HTML
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 调试：保存原始HTML到文件以便分析
            debug_html_path = "debug_page.html"
            with open(debug_html_path, 'w', encoding='utf-8') as f:
                f.write(response.text)
            logger.debug(f"原始HTML已保存到: {debug_html_path}")
            
            # 提取标题
            title = self._extract_title_from_soup(soup)
            
            # 提取作者
            author = self._extract_author_from_soup(soup)
            
            # 提取发布时间
            publish_date = self._extract_publish_date_from_soup(soup)
            
            # 提取文章正文内容
            content_soup = self._extract_content_from_soup(soup)
            
            if not content_soup:
                return {"error": "未找到文章正文内容"}
            
            # 下载图片
            images = self._download_images_from_soup(content_soup, url)
            
            logger.success(f"✅ 成功提取文章内容: {title[:30]}...")
            logger.info(f"📊 内容统计: 段落数={len(content_soup.find_all(['p', 'div']))}, 图片数量={len(images)}")
            
            return {
                'title': title,
                'author': author,
                'publish_date': publish_date,
                'content_soup': content_soup,
                'images': images,
                'url': url
            }
            
        except requests.RequestException as e:
            logger.error(f"网络请求失败: {e}")
            return {"error": f"网络请求失败: {str(e)}"}
        except Exception as e:
            logger.error(f"解析HTML失败: {e}")
            return {"error": f"解析失败: {str(e)}"}
    
    def _extract_title_from_soup(self, soup: BeautifulSoup) -> str:
        """从soup中提取标题"""
        try:
            # 微信文章标题的常见位置
            title_selectors = [
                'h1#activity-name',
                'h2.rich_media_title',
                '.rich_media_title',
                'h1.rich_media_title',
                '.rich_media_title h1',
                '#activity-name',
                'h1',
                'title'
            ]
            
            for selector in title_selectors:
                title_elem = soup.select_one(selector)
                if title_elem and title_elem.get_text(strip=True):
                    title = title_elem.get_text(strip=True)
                    logger.debug(f"通过选择器 {selector} 找到标题: {title[:50]}...")
                    return title
            
            # 如果都没找到，从页面源码中搜索
            page_text = soup.get_text()
            import re
            # 尝试从meta信息中找标题
            meta_title = soup.find('meta', property='og:title')
            if meta_title and meta_title.get('content'):
                title = meta_title.get('content').strip()
                if title:
                    logger.debug(f"从meta标签找到标题: {title[:50]}...")
                    return title
            
            return "未知标题"
            
        except Exception as e:
            logger.debug(f"提取标题失败: {e}")
            return "未知标题"
    
    def _extract_author_from_soup(self, soup: BeautifulSoup) -> str:
        """从soup中提取作者"""
        try:
            # 微信文章作者的常见位置
            author_selectors = [
                '.rich_media_meta_nickname',
                '.profile_nickname',
                '[data-role="nickname"]',
                '.account_nickname'
            ]
            
            for selector in author_selectors:
                author_elem = soup.select_one(selector)
                if author_elem and author_elem.get_text(strip=True):
                    author = author_elem.get_text(strip=True)
                    logger.debug(f"通过选择器 {selector} 找到作者: {author}")
                    return author
            
            return "未知作者"
            
        except Exception as e:
            logger.debug(f"提取作者失败: {e}")
            return "未知作者"
    
    def _extract_publish_date_from_soup(self, soup: BeautifulSoup) -> str:
        """从soup中提取发布时间"""
        try:
            # 发布时间的常见位置
            date_selectors = [
                '#publish_time',
                '.rich_media_meta_text',
                '[data-role="publish-time"]'
            ]
            
            for selector in date_selectors:
                date_elem = soup.select_one(selector)
                if date_elem and date_elem.get_text(strip=True):
                    date_text = date_elem.get_text(strip=True)
                    logger.debug(f"通过选择器 {selector} 找到发布时间: {date_text}")
                    return date_text
            
            # 尝试从脚本中提取
            scripts = soup.find_all('script')
            for script in scripts:
                if script.string and 'publish_time' in script.string:
                    import re
                    match = re.search(r'publish_time["\s]*[=:]["\s]*("[^"]+)"', script.string)
                    if match:
                        return match.group(1)
            
            return "未知时间"
            
        except Exception as e:
            logger.debug(f"提取发布时间失败: {e}")
            return "未知时间"
    
    def _extract_content_from_soup(self, soup: BeautifulSoup) -> BeautifulSoup:
        """从soup中提取文章正文区域"""
        try:
            # 微信文章正文的常见容器
            content_selectors = [
                '#js_content',
                '.rich_media_content',
                '.appmsg_content_text',
                '[data-role="article-content"]',
                '.rich_media_area_primary',
                '.rich_media_wrp'
            ]
            
            for selector in content_selectors:
                content_elem = soup.select_one(selector)
                if content_elem:
                    # 检查内容是否足够多
                    text_content = content_elem.get_text(strip=True)
                    if text_content and len(text_content) > 50:  # 至少50个字符
                        logger.debug(f"通过选择器 {selector} 找到正文内容，长度: {len(text_content)}")
                        return content_elem
            
            # 如果没有找到合适的容器，尝试智能提取
            logger.warning("未找到专用正文容器，尝试智能提取")
            
            # 查找包含较多文本的div
            all_divs = soup.find_all('div')
            best_div = None
            max_text_length = 0
            
            for div in all_divs:
                text = div.get_text(strip=True)
                if len(text) > max_text_length and len(text) > 100:  # 至少100个字符
                    # 排除一些不相关的div
                    div_class = div.get('class', [])
                    div_id = div.get('id', '')
                    if not any(exclude in str(div_class) + div_id for exclude in ['nav', 'header', 'footer', 'sidebar', 'menu']):
                        max_text_length = len(text)
                        best_div = div
            
            if best_div:
                logger.debug(f"智能提取找到内容区域，文本长度: {max_text_length}")
                return best_div
            
            logger.warning("使用整个body作为内容")
            return soup.find('body') or soup
            
        except Exception as e:
            logger.error(f"提取正文内容失败: {e}")
            return None
    
    def _download_images_from_soup(self, content_soup: BeautifulSoup, base_url: str) -> list:
        """从内容中下载图片"""
        try:
            images_info = []
            img_tags = content_soup.find_all('img')
            
            if not img_tags:
                logger.info("📷 未发现图片")
                return images_info
            
            logger.info(f"🖼️ 发现 {len(img_tags)} 张图片，开始下载...")
            
            # 创建图片下载目录
            import urllib.parse
            parsed_url = urllib.parse.urlparse(base_url)
            safe_domain = re.sub(r'[<>:"/\\|?*]', '_', parsed_url.netloc)
            img_dir = os.path.join("output", "images", safe_domain)
            os.makedirs(img_dir, exist_ok=True)
            
            for i, img in enumerate(img_tags):
                try:
                    img_src = img.get('src') or img.get('data-src') or img.get('data-original')
                    if not img_src:
                        continue
                    
                    # 处理相对URL
                    if img_src.startswith('//'):
                        img_src = 'https:' + img_src
                    elif img_src.startswith('/'):
                        img_src = f"https://{parsed_url.netloc}" + img_src
                    elif not img_src.startswith('http'):
                        img_src = urllib.parse.urljoin(base_url, img_src)
                    
                    # 生成本地文件名
                    img_filename = f"img_{i+1:03d}.jpg"
                    img_path = os.path.join(img_dir, img_filename)
                    
                    # 下载图片
                    if self._download_image_requests(img_src, img_path):
                        images_info.append({
                            'url': img_src,
                            'local_path': img_path,
                            'filename': img_filename
                        })
                        logger.debug(f"📷 下载图片成功: {img_filename}")
                    else:
                        images_info.append({
                            'url': img_src,
                            'local_path': None,
                            'filename': img_filename
                        })
                        logger.warning(f"📷 下载图片失败: {img_src}")
                
                except Exception as e:
                    logger.warning(f"处理图片时出错: {e}")
                    continue
            
            logger.success(f"🖼️ 图片下载完成: {len([img for img in images_info if img['local_path']])}/{len(img_tags)}")
            return images_info
            
        except Exception as e:
            logger.error(f"图片下载异常: {e}")
            return []
    
    def _download_image_requests(self, img_url: str, save_path: str) -> bool:
        """使用requests下载图片"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Referer': 'https://mp.weixin.qq.com/'
            }
            
            response = requests.get(img_url, headers=headers, timeout=15, stream=True)
            response.raise_for_status()
            
            with open(save_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            return True
            
        except Exception as e:
            logger.debug(f"下载图片失败 {img_url}: {e}")
            return False
    
    def _process_wechat_content_to_docx(self, doc, content_soup: BeautifulSoup, images: list):
        """
        将微信文章内容转换为Word文档，确保内容完整性
        """
        try:
            # 递归处理所有内容元素，添加深度限制
            self._process_element_to_docx_recursive(doc, content_soup, images, depth=0, max_depth=10)
                    
        except Exception as e:
            logger.warning(f"内容转换异常: {e}")
            # 降级到纯文本处理
            try:
                text = content_soup.get_text()
                if text.strip():
                    doc.add_paragraph(text.strip())
            except:
                doc.add_paragraph("内容提取失败")
    
    def _process_element_to_docx_recursive(self, doc, element, images: list, depth: int = 0, max_depth: int = 10):
        """
        递归处理HTML元素到Word文档（带深度限制）
        """
        try:
            # 防止无限递归
            if depth > max_depth:
                logger.warning(f"递归深度超限 ({depth})，停止处理")
                return
                
            if hasattr(element, 'name'):
                tag_name = element.name.lower() if element.name else None
                
                # 处理不同类型的元素
                if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # 标题处理
                    level = int(tag_name[1])
                    text = element.get_text(strip=True)
                    if text:
                        doc.add_heading(text, level=min(level, 3))
                        
                elif tag_name == 'p':
                    # 段落处理
                    self._process_paragraph_to_docx(doc, element, images)
                    
                elif tag_name == 'div':
                    # div容器，递归处理子元素
                    text = element.get_text(strip=True)
                    if text and len(text) < 200:  # 短文本直接作为段落
                        paragraph = doc.add_paragraph()
                        self._add_formatted_text_to_paragraph(paragraph, element, images)
                    else:
                        # 长内容，递归处理子元素
                        if hasattr(element, 'children'):
                            child_count = 0
                            for child in element.children:
                                child_count += 1
                                if child_count > 50:  # 限制子元素数量
                                    logger.warning("子元素过多，停止处理")
                                    break
                                self._process_element_to_docx_recursive(doc, child, images, depth + 1, max_depth)
                    
                elif tag_name == 'img':
                    # 图片处理
                    self._add_image_to_docx_new(doc, element, images)
                    
                elif tag_name == 'blockquote':
                    # 引用处理
                    text = element.get_text(strip=True)
                    if text:
                        paragraph = doc.add_paragraph(text)
                        # 应用斜体样式
                        for run in paragraph.runs:
                            run.italic = True
                            
                elif tag_name in ['ul', 'ol']:
                    # 列表处理
                    items = element.find_all('li')[:20]  # 限制列表项数量
                    for item in items:
                        text = item.get_text(strip=True)
                        if text:
                            doc.add_paragraph(f"• {text}", style='List Bullet')
                            
                elif tag_name == 'br':
                    # 换行
                    doc.add_paragraph("")
                    
                elif tag_name in ['span', 'strong', 'b', 'em', 'i']:
                    # 内联元素，在父级处理
                    pass
                    
                else:
                    # 其他元素，递归处理子元素
                    if hasattr(element, 'children'):
                        child_count = 0
                        for child in element.children:
                            child_count += 1
                            if child_count > 50:  # 限制子元素数量
                                logger.warning("子元素过多，停止处理")
                                break
                            self._process_element_to_docx_recursive(doc, child, images, depth + 1, max_depth)
            else:
                # 文本节点
                if hasattr(element, 'strip'):
                    text_content = str(element).strip()
                    if text_content and len(text_content) > 0 and len(text_content) < 1000:  # 限制文本长度
                        doc.add_paragraph(text_content)
                    
        except Exception as e:
            logger.debug(f"处理元素异常 (深度{depth}): {e}")
            # 异常时添加简单文本内容
            try:
                if hasattr(element, 'get_text'):
                    text = element.get_text(strip=True)
                    if text and len(text) < 500:
                        doc.add_paragraph(text)
            except:
                pass
    
    def _process_paragraph_to_docx(self, doc, p_element, images: list):
        """处理段落元素"""
        try:
            # 检查段落是否为空
            text_content = p_element.get_text(strip=True)
            if not text_content:
                return
            
            # 创建段落
            paragraph = doc.add_paragraph()
            
            # 添加格式化文本
            self._add_formatted_text_to_paragraph(paragraph, p_element, images)
            
        except Exception as e:
            logger.debug(f"处理段落异常: {e}")
    
    def _add_formatted_text_to_paragraph(self, paragraph, element, images: list):
        """添加格式化文本到段落"""
        try:
            if hasattr(element, 'children'):
                for child in element.children:
                    if hasattr(child, 'name') and child.name:
                        tag_name = child.name.lower()
                        
                        if tag_name == 'img':
                            # 图片插入到段落中
                            try:
                                self._add_image_to_docx_new(paragraph._element.getparent().getparent(), child, images, True)
                            except:
                                pass
                                
                        elif tag_name in ['strong', 'b']:
                            # 加粗文本
                            text = child.get_text()
                            if text:
                                run = paragraph.add_run(text)
                                run.bold = True
                                
                        elif tag_name in ['em', 'i']:
                            # 斜体文本
                            text = child.get_text()
                            if text:
                                run = paragraph.add_run(text)
                                run.italic = True
                                
                        elif tag_name == 'a':
                            # 链接
                            text = child.get_text()
                            href = child.get('href', '')
                            if text:
                                run = paragraph.add_run(f"{text}({href})" if href else text)
                                
                        else:
                            # 其他标签，递归处理
                            self._add_formatted_text_to_paragraph(paragraph, child, images)
                    else:
                        # 文本节点
                        text = str(child).strip()
                        if text:
                            paragraph.add_run(text)
            else:
                # 如果没有children，直接处理文本内容
                text = element.get_text() if hasattr(element, 'get_text') else str(element)
                if text.strip():
                    paragraph.add_run(text.strip())
                            
        except Exception as e:
            logger.debug(f"添加格式化文本异常: {e}")
    
    def _add_image_to_docx_new(self, doc, img_element, images: list, inline: bool = False):
        """添加图片到Word文档（优化版）"""
        try:
            img_src = img_element.get('src') or img_element.get('data-src') or img_element.get('data-original')
            if not img_src:
                return
            
            # 查找对应的本地图片
            local_image = None
            for img_info in images:
                if img_info.get('url') == img_src or img_src in img_info.get('url', ''):
                    local_image = img_info
                    break
            
            if local_image and local_image.get('local_path'):
                local_path = local_image['local_path']
                
                # 验证文件是否存在且有效
                if not os.path.exists(local_path):
                    logger.debug(f"图片文件不存在: {local_path}")
                    if not inline:
                        doc.add_paragraph(f"[图片: {img_src}]")
                    return
                
                # 检查文件大小
                try:
                    file_size = os.path.getsize(local_path)
                    if file_size == 0:
                        logger.debug(f"图片文件为空: {local_path}")
                        if not inline:
                            doc.add_paragraph(f"[图片: {img_src}]")
                        return
                    elif file_size > 5 * 1024 * 1024:  # 5MB限制
                        logger.debug(f"图片文件过大: {file_size} bytes")
                        if not inline:
                            doc.add_paragraph(f"[图片过大: {img_src}]")
                        return
                except OSError:
                    logger.debug(f"无法获取图片文件大小: {local_path}")
                    if not inline:
                        doc.add_paragraph(f"[图片: {img_src}]")
                    return
                
                # 尝试插入图片
                try:
                    # 验证图片格式
                    valid_extensions = ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp')
                    if not any(local_path.lower().endswith(ext) for ext in valid_extensions):
                        # 尝试重命名为.jpg
                        new_path = local_path.rsplit('.', 1)[0] + '.jpg'
                        if os.path.exists(local_path):
                            try:
                                os.rename(local_path, new_path)
                                local_path = new_path
                            except:
                                pass
                    
                    # 设置图片大小
                    if inline:
                        # 内联图片，较小尺寸
                        doc.add_picture(local_path, width=Inches(3))
                    else:
                        # 独立段落图片，较大尺寸
                        doc.add_picture(local_path, width=Inches(4.5))
                    
                    logger.debug(f"✅ 图片插入成功: {os.path.basename(local_path)}")
                    
                except Exception as img_error:
                    logger.debug(f"图片插入失败 {local_path}: {img_error}")
                    # 降级处理：添加图片链接文本
                    if not inline:
                        doc.add_paragraph(f"[图片链接: {img_src}]")
            else:
                # 没有本地图片文件，添加链接文本
                if not inline:
                    doc.add_paragraph(f"[图片链接: {img_src}]")
                
        except Exception as e:
            logger.debug(f"图片处理异常: {e}")
            # 静默失败，不影响整体文档生成
    
    def _process_html_to_docx(self, doc, html_content: str, images: list):
        """
        将HTML内容转换为Word文档，确保内容完整性
        """
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 递归处理每个元素
            for element in soup.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li', 'img', 'br']):
                self._process_element_to_docx(doc, element, images)
            
            # 如果没有找到标准元素，处理所有文本
            if not soup.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                all_text = soup.get_text(strip=True)
                if all_text:
                    doc.add_paragraph(all_text)
                    
        except Exception as e:
            logger.warning(f"HTML转换异常: {e}")
            # 降级到纯文本处理
            try:
                soup = BeautifulSoup(html_content, 'html.parser')
                text = soup.get_text()
                doc.add_paragraph(text)
            except:
                doc.add_paragraph("内容提取失败")
    
    def _process_element_to_docx(self, doc, element, images: list):
        """
        处理单个HTML元素到Word文档
        """
        try:
            tag_name = element.name.lower()
            
            # 标题处理
            if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag_name[1])
                text = element.get_text(strip=True)
                if text:
                    doc.add_heading(text, level=min(level, 3))
            
            # 段落处理
            elif tag_name in ['p', 'div']:
                text = element.get_text(strip=True)
                if text:
                    paragraph = doc.add_paragraph(text)
                    
                    # 检查是否包含图片
                    imgs = element.find_all('img')
                    for img in imgs:
                        self._add_image_to_docx(doc, img, images)
            
            # 列表处理
            elif tag_name in ['ul', 'ol']:
                items = element.find_all('li')
                for item in items:
                    text = item.get_text(strip=True)
                    if text:
                        doc.add_paragraph(f"• {text}", style='List Bullet')
            
            # 图片处理
            elif tag_name == 'img':
                self._add_image_to_docx(doc, element, images)
            
            # 换行处理
            elif tag_name == 'br':
                doc.add_paragraph("")
                
        except Exception as e:
            logger.debug(f"元素处理异常: {e}")
    
    def _add_image_to_docx(self, doc, img_element, images: list):
        """
        添加图片到Word文档
        """
        try:
            img_src = img_element.get('src') or img_element.get('data-src') or img_element.get('data-original')
            if not img_src:
                return
            
            # 查找对应的本地图片
            local_image = None
            for img_info in images:
                if img_info.get('url') == img_src or img_src in img_info.get('url', ''):
                    local_image = img_info
                    break
            
            if local_image and local_image.get('local_path') and os.path.exists(local_image['local_path']):
                try:
                    doc.add_picture(local_image['local_path'], width=Inches(4))
                    logger.debug(f"添加图片成功: {local_image['local_path']}")
                except Exception as e:
                    logger.warning(f"图片添加失败: {e}")
                    doc.add_paragraph(f"[图片: {img_src}]")
            else:
                doc.add_paragraph(f"[图片链接: {img_src}]")
                
        except Exception as e:
            logger.debug(f"图片处理异常: {e}")
    
    def _handle_captcha_if_present(self) -> bool:
        """处理验证码（如果存在）"""
        try:
            logger.info("🔍 检查是否存在验证码...")
            
            # 等待页面稳定
            time.sleep(3)
            
            # 检查常见的验证码元素
            captcha_selectors = [
                "iframe[src*='captcha']",
                "div[class*='captcha']",
                "div[id*='captcha']",
                ".captcha-container",
                "#captcha",
                "iframe[src*='verify']",
                "div[class*='verify']",
                "canvas[id*='captcha']",
                "img[src*='captcha']"
            ]
            
            captcha_found = False
            for selector in captcha_selectors:
                try:
                    elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements:
                        logger.warning(f"🎯 发现验证码元素: {selector}")
                        captcha_found = True
                        break
                except:
                    continue
            
            if captcha_found:
                logger.warning("⚠️ 检测到验证码，切换到非headless模式以便手动处理")
                
                # 如果是headless模式，提示用户
                if self.headless:
                    logger.error("❌ 检测到验证码但当前为headless模式，请使用GUI模式运行程序")
                    return False
                
                # 等待用户处理验证码
                logger.info("🖱️ 请在浏览器中完成验证码验证...")
                logger.info("⏳ 程序将等待60秒供您完成验证...")
                
                # 等待验证码消失或页面跳转
                for i in range(60):
                    time.sleep(1)
                    try:
                        # 检查验证码是否还存在
                        still_has_captcha = False
                        for selector in captcha_selectors:
                            try:
                                elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                                if elements and elements[0].is_displayed():
                                    still_has_captcha = True
                                    break
                            except:
                                continue
                        
                        if not still_has_captcha:
                            logger.info("✅ 验证码已完成，继续处理...")
                            return True
                            
                        # 检查是否已经跳转到文章页面
                        if "mp.weixin.qq.com/s" in self.driver.current_url:
                            article_title = self.driver.find_elements(By.CSS_SELECTOR, "#activity-name, .rich_media_title")
                            if article_title:
                                logger.info("✅ 已成功跳转到文章页面")
                                return True
                                
                    except Exception as e:
                        logger.debug(f"验证码检查异常: {e}")
                        continue
                
                logger.error("⏰ 验证码处理超时，请重试")
                return False
            else:
                logger.info("✅ 未检测到验证码")
                return True
                
        except Exception as e:
            logger.error(f"验证码处理异常: {e}")
            return True  # 继续执行，可能没有验证码

    def _wait_for_basic_page_load(self) -> bool:
        """
        等待页面基础内容加载 - 优化版本
        """
        try:
            logger.info("等待页面基础内容加载...")
            
            # 减少等待时间到最多5秒
            wait = WebDriverWait(self.driver, 5)
            
            # 等待页面基本元素加载
            wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            
            # 等待文章内容加载（减少等待时间）
            try:
                wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 
                    "#js_content, .rich_media_content, .article-content, .content")))
            except:
                logger.debug("未找到常见文章内容容器，使用通用检查")
            
            # 短暂等待DOM稳定
            time.sleep(0.5)
            
            logger.success("页面基础内容加载完成")
            return True
            
        except TimeoutException:
            logger.warning("页面加载超时，但继续处理")
            return True
        except Exception as e:
            logger.error(f"等待页面加载时出错: {e}")
            return False
    
    def _human_like_scroll_and_load(self, target_url: str = None) -> bool:
        """
        人类式滚动加载内容 - 优化版本（减少重复滚动）
        """
        try:
            # 获取页面总高度
            total_height = self.driver.execute_script("return document.body.scrollHeight")
            logger.info(f"开始智能滚动加载，页面高度: {total_height}px")
            
            # 如果页面很短，直接返回
            if total_height <= 1000:
                logger.debug("页面较短，无需滚动")
                return True
            
            # 根据页面高度动态调整滚动策略
            if total_height <= 3000:
                # 短页面：快速滚动
                pixels_per_step = 800
                scroll_delay = 1.0
                check_frequency = 3
            elif total_height <= 8000:
                # 中等页面：平衡滚动
                pixels_per_step = 1000
                scroll_delay = 1.2
                check_frequency = 4
            else:
                # 长页面：大步滚动
                pixels_per_step = 1500
                scroll_delay = 1.5
                check_frequency = 5
            
            # 计算滚动步数
            scroll_positions = max(3, (total_height // pixels_per_step) + 1)
            scroll_positions = min(scroll_positions, 8)  # 最多8个位置，避免过度滚动
            
            logger.info(f"滚动策略: {scroll_positions}步, 每步{pixels_per_step}px, 延迟{scroll_delay}s")
            
            # 执行滚动
            for i in range(scroll_positions):
                # 计算滚动位置
                if scroll_positions == 1:
                    scroll_to = total_height
                else:
                    progress = i / (scroll_positions - 1)
                    scroll_to = int(total_height * progress)
                
                logger.debug(f"滚动到位置 {i+1}/{scroll_positions}: {scroll_to}px")
                
                # 平滑滚动到目标位置
                self.driver.execute_script(f"window.scrollTo({{top: {scroll_to}, behavior: 'smooth'}});")
                
                # 等待内容加载
                time.sleep(scroll_delay)
                
                # 减少图片检查频率
                if i % check_frequency == 0:
                    self._trigger_image_loading()
                    time.sleep(0.5)
            
            # 最终处理：快速滚动到底部
            logger.info("最终处理...")
            self.driver.execute_script("window.scrollTo({top: document.body.scrollHeight, behavior: 'smooth'});")
            time.sleep(1.5)
            
            # 最后一次图片检查
            self._trigger_image_loading()
            time.sleep(1.0)
            
            # 回到顶部
            self.driver.execute_script("window.scrollTo({top: 0, behavior: 'smooth'});")
            time.sleep(0.5)
            
            logger.success("智能滚动完成")
            return True
            
        except Exception as e:
            logger.warning(f"滚动过程异常: {e}")
            return True  # 即使出错也继续，不影响内容提取
    
    def _trigger_image_loading(self):
        """触发图片懒加载"""
        try:
            # 触发各种懒加载机制
            self.driver.execute_script("""
                // 触发懒加载图片
                var images = document.querySelectorAll('img[data-src], img[data-lazy-src], img[data-original]');
                images.forEach(function(img) {
                    if (img.dataset.src) img.src = img.dataset.src;
                    if (img.dataset.lazySrc) img.src = img.dataset.lazySrc;
                    if (img.dataset.original) img.src = img.dataset.original;
                    // 触发onload事件
                    img.onload = function() { console.log('Image loaded:', this.src); };
                });
                
                // 触发滚动事件，可能激活某些懒加载脚本
                window.dispatchEvent(new Event('scroll'));
                window.dispatchEvent(new Event('resize'));
            """)
            
            logger.debug("触发图片懒加载")
            
        except Exception as e:
            logger.debug(f"触发图片懒加载失败: {e}")
    
    def _final_image_check(self):
        """最终图片检查和加载"""
        try:
            logger.info("执行最终快速检查...")
            
            # 最后一次触发所有可能的懒加载
            self._trigger_image_loading()
            
            # 等待图片加载
            time.sleep(3.0)
            
            # 检查图片加载状态
            image_stats = self.driver.execute_script("""
                var images = document.querySelectorAll('img');
                var total = images.length;
                var loaded = 0;
                var broken = 0;
                var unloaded = 0;
                
                images.forEach(function(img) {
                    if (img.complete) {
                        if (img.naturalHeight !== 0) {
                            loaded++;
                        } else {
                            broken++;
                        }
                    } else {
                        unloaded++;
                    }
                });
                
                return {total: total, loaded: loaded, broken: broken, unloaded: unloaded};
            """)
            
            logger.info(f"图片状态: 总数={image_stats['total']}, 已加载={image_stats['loaded']}, 损坏={image_stats['broken']}, 未加载={image_stats['unloaded']}")
            
            # 如果还有未加载的图片，再等待一会
            if image_stats['unloaded'] > 0:
                logger.warning(f"还有 {image_stats['unloaded']} 张图片未加载，额外等待...")
                time.sleep(2.0)
            
            logger.success("最终检查完成")
            
        except Exception as e:
            logger.debug(f"最终图片检查失败: {e}")
    
    def _download_image(self, img_url: str, save_path: str) -> bool:
        """下载图片"""
        try:
            # 处理相对路径
            if img_url.startswith('//'):
                img_url = 'https:' + img_url
            elif img_url.startswith('/'):
                img_url = 'https://mp.weixin.qq.com' + img_url
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Referer': 'https://mp.weixin.qq.com/'
            }
            
            response = requests.get(img_url, headers=headers, timeout=10)
            response.raise_for_status()
            
            with open(save_path, 'wb') as f:
                f.write(response.content)
            
            return True
            
        except Exception as e:
            logger.debug(f"下载图片失败 {img_url}: {e}")
            return False
    
    def _get_article_title(self, soup: BeautifulSoup) -> str:
        """从BeautifulSoup对象中提取文章标题"""
        try:
            # 尝试多种选择器
            title_selectors = [
                '#activity-name',
                '.rich_media_title',
                'h1',
                '.article-title',
                '.title'
            ]
            
            for selector in title_selectors:
                title_elem = soup.select_one(selector)
                if title_elem and title_elem.get_text(strip=True):
                    title = title_elem.get_text(strip=True)
                    # 清理标题
                    title = re.sub(r'\s+', ' ', title)
                    return title[:50]
                    
        except Exception as e:
            logger.debug(f"提取标题失败: {e}")
        
        return "unknown_article"
    
    def _extract_article_content(self) -> dict:
        """
        提取文章内容信息 - 快速版本
        """
        try:
            article_info = {}
            
            # 快速获取标题
            title_selectors = [
                '#activity-name',
                '.rich_media_title', 
                'h1',
                '.title',
                '[data-role="title"]'
            ]
            
            title = None
            for selector in title_selectors:
                try:
                    title_elem = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if title_elem and title_elem.text.strip():
                        title = title_elem.text.strip()
                        break
                except:
                    continue
            
            article_info['title'] = title or "未知标题"
            
            # 快速获取作者
            author_selectors = [
                '#js_name',
                '.profile_nickname',
                '.author',
                '[data-role="author"]'
            ]
            
            author = None
            for selector in author_selectors:
                try:
                    author_elem = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if author_elem and author_elem.text.strip():
                        author = author_elem.text.strip()
                        logger.debug(f"通过选择器 {selector} 找到作者: {author}")
                        break
                except:
                    continue
            
            article_info['author'] = author or "未知作者"
            logger.info(f"最终识别的公众号/作者: {article_info['author']}")
            
            # 添加URL
            article_info['url'] = self.driver.current_url
            
            logger.success(f"成功提取文章信息: {article_info['title'][:30]}...")
            return article_info
            
        except Exception as e:
            logger.error(f"提取文章内容失败: {e}")
            return {"error": f"内容提取失败: {str(e)}"}
    
    def process_urls(self, urls: List[str]) -> List[Dict[str, Any]]:
        """批量处理URL列表"""
        try:
            logger.info(f"开始处理 {len(urls)} 个URL...")
            
            articles = []
            success_count = 0
            
            from tqdm import tqdm
            for i, url in enumerate(tqdm(urls, desc="处理URL")):
                try:
                    logger.info(f"处理第 {i+1}/{len(urls)} 个URL...")
                    
                    # 验证URL格式
                    if not self._is_valid_wechat_url(url):
                        logger.warning(f"无效的微信文章URL: {url}")
                        articles.append({
                            'title': f"无效URL_{i+1}",
                            'author': "未知",
                            'publish_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            'content': "无效的URL格式",
                            'url': url
                        })
                        continue
                    
                    # 提取文章信息
                    article_info = self.extract_article_info(url)
                    
                    if 'error' not in article_info:
                        articles.append(article_info)
                        if len(article_info.get('content', '').strip()) > 100:
                            success_count += 1
                    else:
                        logger.error(f"处理URL失败: {article_info.get('error')}")
                        articles.append({
                            'title': f"错误_文章_{i+1}",
                            'author': "未知",
                            'publish_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            'content': f"处理失败: {article_info.get('error')}",
                            'url': url
                        })
                    
                    # 添加延迟避免过于频繁的请求
                    if i < len(urls) - 1:  # 不是最后一个URL
                        delay = random.uniform(2, 5)
                        logger.debug(f"延迟 {delay:.2f} 秒...")
                        time.sleep(delay)
                        
                except Exception as e:
                    logger.error(f"处理第 {i+1} 个URL时出错: {e}")
                    articles.append({
                        'title': f"异常_文章_{i+1}",
                        'author': "未知",
                        'publish_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'content': f"处理异常: {str(e)}",
                        'url': url
                    })
            
            logger.success(f"处理完成！成功: {success_count}/{len(articles)}")
            return articles
            
        except Exception as e:
            logger.error(f"批量处理失败: {e}")
            return []
    
    def process_urls_as_pdf(self, urls: List[str], output_dir: str = "output/pdf") -> List[str]:
        """批量处理URL并保存为PDF格式"""
        try:
            logger.info(f"开始处理 {len(urls)} 个URL并保存为PDF...")
            
            os.makedirs(output_dir, exist_ok=True)
            saved_files = []
            success_count = 0
            
            from tqdm import tqdm
            for i, url in enumerate(tqdm(urls, desc="生成PDF")):
                try:
                    logger.info(f"处理第 {i+1}/{len(urls)} 个URL...")
                    
                    # 验证URL格式
                    if not self._is_valid_wechat_url(url):
                        logger.warning(f"无效的微信文章URL: {url}")
                        continue
                    
                    # 先获取文章信息用于命名
                    article_info = self.extract_article_info(url)
                    
                    if 'error' in article_info:
                        logger.error(f"获取文章信息失败: {article_info.get('error')}")
                        continue
                    
                    # 生成PDF文件名
                    title = article_info.get('title', f'文章_{i+1}')
                    safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:50]
                    pdf_filename = f"{safe_title}.pdf"
                    pdf_path = os.path.join(output_dir, pdf_filename)
                    
                    # 处理文件名冲突
                    counter = 1
                    original_path = pdf_path
                    while os.path.exists(pdf_path):
                        name_without_ext = safe_title
                        pdf_filename = f"{name_without_ext}_{counter}.pdf"
                        pdf_path = os.path.join(output_dir, pdf_filename)
                        counter += 1
                    
                    # 保存为PDF
                    if self.save_as_pdf(url, pdf_path):
                        saved_files.append(pdf_path)
                        success_count += 1
                        logger.success(f"PDF已保存: {pdf_filename}")
                    
                    # 添加延迟
                    if i < len(urls) - 1:
                        delay = random.uniform(2, 5)
                        logger.debug(f"延迟 {delay:.2f} 秒...")
                        time.sleep(delay)
                        
                except Exception as e:
                    logger.error(f"处理第 {i+1} 个URL时出错: {e}")
            
            logger.success(f"PDF生成完成！成功: {success_count}/{len(urls)}")
            return saved_files
            
        except Exception as e:
            logger.error(f"批量PDF生成失败: {e}")
            return []
    
    def process_urls_as_docx(self, urls: List[str], output_dir: str = "output/docx") -> List[str]:
        """批量处理URL并保存为Word文档格式"""
        try:
            logger.info(f"开始处理 {len(urls)} 个URL并保存为Word文档...")
            
            os.makedirs(output_dir, exist_ok=True)
            saved_files = []
            success_count = 0
            
            from tqdm import tqdm
            for i, url in enumerate(tqdm(urls, desc="生成Word文档")):
                try:
                    logger.info(f"处理第 {i+1}/{len(urls)} 个URL...")
                    
                    # 验证URL格式
                    if not self._is_valid_wechat_url(url):
                        logger.warning(f"无效的微信文章URL: {url}")
                        continue
                    
                    # 先获取文章信息用于命名
                    article_info = self.extract_article_info(url)
                    
                    if 'error' in article_info:
                        logger.error(f"获取文章信息失败: {article_info.get('error')}")
                        continue
                    
                    # 生成Word文件名
                    title = article_info.get('title', f'文章_{i+1}')
                    safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:50]
                    docx_filename = f"{safe_title}.docx"
                    docx_path = os.path.join(output_dir, docx_filename)
                    
                    # 处理文件名冲突
                    counter = 1
                    original_path = docx_path
                    while os.path.exists(docx_path):
                        name_without_ext = safe_title
                        docx_filename = f"{name_without_ext}_{counter}.docx"
                        docx_path = os.path.join(output_dir, docx_filename)
                        counter += 1
                    
                    # 保存为Word文档
                    if self.save_as_docx(url, docx_path):
                        saved_files.append(docx_path)
                        success_count += 1
                        logger.success(f"Word文档已保存: {docx_filename}")
                    
                    # 添加延迟
                    if i < len(urls) - 1:
                        delay = random.uniform(2, 5)
                        logger.debug(f"延迟 {delay:.2f} 秒...")
                        time.sleep(delay)
                        
                except Exception as e:
                    logger.error(f"处理第 {i+1} 个URL时出错: {e}")
            
            logger.success(f"Word文档生成完成！成功: {success_count}/{len(urls)}")
            return saved_files
            
        except Exception as e:
            logger.error(f"批量Word文档生成失败: {e}")
            return []
    
    def _is_valid_wechat_url(self, url: str) -> bool:
        """验证是否是有效的微信文章URL"""
        try:
            # 检查基本格式
            if not url or not url.startswith('http'):
                return False
            
            # 检查是否是微信文章URL
            if 'mp.weixin.qq.com/s' not in url:
                return False
            
            return True
            
        except Exception:
            return False
    
    def cleanup(self):
        """清理资源"""
        try:
            if self.driver:
                self.driver.quit()
                logger.debug("浏览器已关闭")
        except Exception as e:
            logger.debug(f"清理资源时出错: {e}")
        finally:
            self.driver = None
    
    def save_complete_html(self, url: str, output_path: str) -> bool:
        """保存完整的HTML文件，包括图片和样式（修复版）"""
        try:
            logger.info(f"正在保存完整HTML: {url}")
            
            # 使用已有的文章提取方法，避免重复滚动
            article_data = self._extract_wechat_article_with_selenium(url)
            
            if not article_data or 'error' in article_data:
                logger.error("无法获取文章内容")
                return False
            
            # 创建输出目录
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 获取文章信息
            title = article_data.get('title', '未知标题')
            author = article_data.get('author', '未知作者')
            publish_date = article_data.get('publish_date', '未知时间')
            content_soup = article_data.get('content_soup')
            images = article_data.get('images', [])
            
            if not content_soup:
                logger.error("没有找到文章内容")
                return False
            
            # 创建完整的HTML文档
            html_doc = BeautifulSoup('<!DOCTYPE html><html><head></head><body></body></html>', 'html.parser')
            
            # 设置HTML头部
            html_doc.head.append(html_doc.new_tag('meta', charset='utf-8'))
            html_doc.head.append(html_doc.new_tag('meta', attrs={'name': 'viewport', 'content': 'width=device-width, initial-scale=1.0'}))
            
            title_tag = html_doc.new_tag('title')
            title_tag.string = title
            html_doc.head.append(title_tag)
            
            # 添加CSS样式
            style_tag = html_doc.new_tag('style')
            style_tag.string = """
                body { 
                    max-width: 800px; 
                    margin: 0 auto; 
                    padding: 20px; 
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif; 
                    line-height: 1.6;
                    color: #333;
                }
                .article-header {
                    text-align: center;
                    margin-bottom: 30px;
                    padding-bottom: 20px;
                    border-bottom: 1px solid #eee;
                }
                .article-title {
                    font-size: 24px;
                    font-weight: bold;
                    margin-bottom: 10px;
                    color: #2c3e50;
                }
                .article-meta {
                    color: #666;
                    font-size: 14px;
                }
                .article-content {
                    font-size: 16px;
                    line-height: 1.8;
                }
                img { 
                    max-width: 100%; 
                    height: auto; 
                    display: block;
                    margin: 15px auto;
                    border-radius: 4px;
                    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
                }
                p { margin: 15px 0; }
                blockquote {
                    border-left: 4px solid #3498db;
                    margin: 20px 0;
                    padding-left: 15px;
                    color: #666;
                    font-style: italic;
                }
                .image-placeholder {
                    background: #f8f9fa;
                    border: 2px dashed #dee2e6;
                    padding: 20px;
                    text-align: center;
                    color: #6c757d;
                    margin: 15px 0;
                    border-radius: 4px;
                }
            """
            html_doc.head.append(style_tag)
            
            # 创建文章头部
            header_div = html_doc.new_tag('div', class_='article-header')
            
            title_h1 = html_doc.new_tag('h1', class_='article-title')
            title_h1.string = title
            header_div.append(title_h1)
            
            meta_div = html_doc.new_tag('div', class_='article-meta')
            meta_div.string = f"作者: {author} | 发布时间: {publish_date} | 原文链接: {url}"
            header_div.append(meta_div)
            
            html_doc.body.append(header_div)
            
            # 创建文章内容容器
            content_div = html_doc.new_tag('div', class_='article-content')
            
            # 处理图片并替换为本地路径
            base_dir = os.path.dirname(output_path)
            safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:20]
            images_dir = os.path.join(base_dir, f"images_{safe_title}")
            os.makedirs(images_dir, exist_ok=True)
            
            # 替换图片链接
            img_count = 0
            for img_tag in content_soup.find_all('img'):
                img_src = img_tag.get('src') or img_tag.get('data-src') or img_tag.get('data-original')
                if img_src:
                    # 查找对应的本地图片
                    local_img_path = None
                    for img_info in images:
                        if img_info.get('url') == img_src or img_src in img_info.get('url', ''):
                            if img_info.get('local_path') and os.path.exists(img_info['local_path']):
                                # 复制图片到HTML目录
                                img_filename = f"img_{img_count:03d}.jpg"
                                new_img_path = os.path.join(images_dir, img_filename)
                                try:
                                    import shutil
                                    shutil.copy2(img_info['local_path'], new_img_path)
                                    local_img_path = f"images_{safe_title}/{img_filename}"
                                    img_count += 1
                                except Exception as e:
                                    logger.debug(f"复制图片失败: {e}")
                                break
                    
                    if local_img_path:
                        img_tag['src'] = local_img_path
                        # 清理其他属性
                        for attr in ['data-src', 'data-original', 'data-lazy-src']:
                            if img_tag.get(attr):
                                del img_tag[attr]
                    else:
                        # 创建图片占位符
                        placeholder_div = html_doc.new_tag('div', class_='image-placeholder')
                        placeholder_div.string = f"[图片: {img_src}]"
                        img_tag.replace_with(placeholder_div)
            
            # 将处理后的内容添加到HTML文档
            for element in content_soup.children:
                if hasattr(element, 'name'):
                    content_div.append(element)
            
            html_doc.body.append(content_div)
            
            # 保存HTML文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(str(html_doc.prettify()))
            
            logger.success(f"完整HTML已保存: {output_path}")
            logger.info(f"📊 统计信息: 图片={img_count}张, 图片目录={images_dir}")
            return True
            
        except Exception as e:
            logger.error(f"保存完整HTML失败: {e}")
            import traceback
            logger.debug(f"详细错误信息: {traceback.format_exc()}")
            return False
    
    def extract_full_article_content(self, url: str, download_images: bool = True) -> dict:
        """
        提取完整的文章内容，包括HTML内容和图片
        
        Args:
            url: 文章URL
            download_images: 是否下载图片到本地
            
        Returns:
            包含完整内容的字典
        """
        try:
            if not self.driver:
                self.driver = self.setup_browser(self.headless)
                if not self.driver:
                    return {"error": "浏览器初始化失败"}

            logger.info(f"🚀 开始提取完整文章内容: {url}")
            
            # 访问页面
            self.driver.get(url)
            
            # 等待页面加载
            try:
                wait = WebDriverWait(self.driver, 10)
                wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            except TimeoutException:
                logger.warning("页面加载超时，尝试继续提取")
            
            # 等待内容加载
            self._wait_for_basic_page_load()
            
            # 滚动加载完整内容
            self._human_like_scroll_and_load()
            
            # 获取基本信息
            basic_info = self._extract_article_content()
            if 'error' in basic_info:
                return basic_info
            
            # 提取完整HTML内容
            html_content = self._extract_html_content()
            
            # 提取发布时间
            publish_date = self._extract_publish_date()
            
            # 处理图片
            images_info = []
            if download_images:
                images_info = self._extract_and_download_images()
            
            # 构建完整的文章信息
            full_article = {
                'title': basic_info.get('title', '未知标题'),
                'author': basic_info.get('author', '未知作者'),
                'url': basic_info.get('url', url),
                'publish_date': publish_date,
                'html_content': html_content,
                'text_content': self._html_to_text(html_content),
                'images': images_info,
                'extraction_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            logger.success(f"✅ 成功提取完整文章内容: {full_article['title'][:30]}...")
            logger.info(f"📊 内容统计: 文字长度={len(full_article['text_content'])}, 图片数量={len(images_info)}")
            
            return full_article
            
        except Exception as e:
            logger.error(f"❌ 提取完整文章内容失败: {e}")
            return {"error": f"提取失败: {str(e)}"}
    
    def _extract_html_content(self) -> str:
        """提取文章的HTML内容"""
        try:
            # 微信文章内容的常见选择器
            content_selectors = [
                '#js_content',
                '.rich_media_content', 
                '.appmsg_content_text',
                '.article_content',
                '[data-role="article-content"]'
            ]
            
            for selector in content_selectors:
                try:
                    content_elem = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if content_elem:
                        # 获取元素的HTML内容
                        html_content = content_elem.get_attribute('innerHTML')
                        if html_content and len(html_content.strip()) > 100:
                            logger.debug(f"通过选择器 {selector} 找到内容，长度: {len(html_content)}")
                            return html_content
                except:
                    continue
            
            # 如果没有找到特定容器，尝试获取整个body
            logger.warning("未找到专用内容容器，尝试提取body内容")
            body_elem = self.driver.find_element(By.TAG_NAME, "body")
            return body_elem.get_attribute('innerHTML') or ""
            
        except Exception as e:
            logger.error(f"提取HTML内容失败: {e}")
            return ""
    
    def _extract_publish_date(self) -> str:
        """提取文章发布时间"""
        try:
            # 发布时间的常见选择器
            date_selectors = [
                '#publish_time',
                '.rich_media_meta_text',
                '.publish_time',
                '[data-role="publish-time"]',
                '.time'
            ]
            
            for selector in date_selectors:
                try:
                    date_elem = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if date_elem and date_elem.text.strip():
                        date_text = date_elem.text.strip()
                        logger.debug(f"通过选择器 {selector} 找到发布时间: {date_text}")
                        return date_text
                except:
                    continue
            
            # 如果没有找到，尝试从页面源码中提取
            page_source = self.driver.page_source
            import re
            
            # 尝试匹配常见的时间格式
            date_patterns = [
                r'publish_time["\s]*[=:]["\s]*("[^"]+)"',
                r'(\d{4}年\d{1,2}月\d{1,2}日)',
                r'(\d{4}-\d{1,2}-\d{1,2})',
                r'(\d{1,2}月\d{1,2}日)'
            ]
            
            for pattern in date_patterns:
                match = re.search(pattern, page_source)
                if match:
                    date_text = match.group(1)
                    logger.debug(f"通过正则表达式找到发布时间: {date_text}")
                    return date_text
            
            return datetime.now().strftime('%Y-%m-%d')
            
        except Exception as e:
            logger.debug(f"提取发布时间失败: {e}")
            return datetime.now().strftime('%Y-%m-%d')
    
    def _extract_and_download_images(self) -> List[Dict[str, str]]:
        """提取并下载文章中的图片"""
        try:
            images_info = []
            
            # 查找所有图片
            img_elements = self.driver.find_elements(By.TAG_NAME, "img")
            
            logger.info(f"🖼️ 发现 {len(img_elements)} 张图片")
            
            for i, img_elem in enumerate(img_elements):
                try:
                    # 获取图片URL
                    img_src = img_elem.get_attribute('src') or img_elem.get_attribute('data-src')
                    
                    if not img_src or img_src.startswith('data:'):
                        continue  # 跳过base64图片和无效图片
                    
                    # 获取图片的alt属性作为描述
                    img_alt = img_elem.get_attribute('alt') or f"图片_{i+1}"
                    
                    # 生成本地文件名
                    img_filename = f"image_{i+1:03d}.jpg"
                    
                    # 创建图片保存目录
                    images_dir = os.path.join("output", "images", 
                                            datetime.now().strftime('%Y%m%d'))
                    os.makedirs(images_dir, exist_ok=True)
                    
                    # 本地保存路径
                    local_path = os.path.join(images_dir, img_filename)
                    
                    # 下载图片
                    if self._download_image(img_src, local_path):
                        images_info.append({
                            'original_url': img_src,
                            'local_path': local_path,
                            'filename': img_filename,
                            'alt': img_alt,
                            'index': i + 1
                        })
                        logger.debug(f"✅ 图片下载成功: {img_filename}")
                    else:
                        logger.warning(f"❌ 图片下载失败: {img_src}")
                
                except Exception as e:
                    logger.debug(f"处理图片 {i+1} 时出错: {e}")
                    continue
            
            logger.success(f"🖼️ 图片下载完成: {len(images_info)}/{len(img_elements)}")
            return images_info
            
        except Exception as e:
            logger.error(f"提取和下载图片失败: {e}")
            return []
    
    def _html_to_text(self, html_content: str) -> str:
        """将HTML内容转换为纯文本"""
        try:
            from bs4 import BeautifulSoup
            
            # 解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除脚本和样式
            for element in soup(["script", "style"]):
                element.decompose()
            
            # 获取纯文本
            text = soup.get_text(separator='\n', strip=True)
            
            # 清理多余的空行
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            clean_text = '\n'.join(lines)
            
            return clean_text
            
        except Exception as e:
            logger.error(f"HTML转文本失败: {e}")
            return ""
    
    def save_as_markdown(self, url: str, output_path: str) -> bool:
        """保存URL为Markdown格式（优化版）"""
        try:
            logger.info(f"正在保存Markdown文档: {url}")
            
            # 使用已有的文章提取方法，避免重复滚动
            article_data = self._extract_wechat_article_with_selenium(url)
            
            if not article_data or 'error' in article_data:
                logger.error("无法获取文章内容")
                return False
            
            # 生成Markdown内容
            markdown_content = []
            
            # 添加标题
            title = article_data.get('title', '微信文章')
            markdown_content.append(f"# {title}\n")
            
            # 添加元信息
            author = article_data.get('author', '未知')
            publish_date = article_data.get('publish_date', '未知')
            markdown_content.append(f"**作者**: {author}")
            markdown_content.append(f"**发布时间**: {publish_date}")
            markdown_content.append(f"**原文链接**: {url}")
            markdown_content.append("\n---\n")
            
            # 转换HTML内容为Markdown
            content_soup = article_data.get('content_soup')
            if content_soup:
                markdown_text = self._convert_soup_to_markdown(content_soup)
                markdown_content.append(markdown_text)
            else:
                markdown_content.append("无法提取文章内容")
            
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 保存Markdown文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_content))
            
            logger.success(f"Markdown文档保存成功: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"保存Markdown文档失败: {e}")
            return False
    
    def save_as_json(self, url: str, output_path: str) -> bool:
        """保存URL为JSON格式（优化版）"""
        try:
            logger.info(f"正在保存JSON数据: {url}")
            
            # 使用已有的文章提取方法，避免重复滚动
            article_data = self._extract_wechat_article_with_selenium(url)
            
            if not article_data or 'error' in article_data:
                logger.error("无法获取文章内容")
                return False
            
            # 构建JSON数据
            json_data = {
                "title": article_data.get('title', ''),
                "author": article_data.get('author', ''),
                "publish_date": article_data.get('publish_date', ''),
                "url": url,
                "extraction_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "content": {
                    "html": str(article_data.get('content_soup', '')),
                    "text": article_data.get('content_soup', BeautifulSoup('', 'html.parser')).get_text(strip=True) if article_data.get('content_soup') else ''
                },
                "images": article_data.get('images', []),
                "metadata": {
                    "total_images": len(article_data.get('images', [])),
                    "content_length": len(article_data.get('content_soup', BeautifulSoup('', 'html.parser')).get_text(strip=True)) if article_data.get('content_soup') else 0,
                    "extraction_method": "selenium"
                }
            }
            
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 保存JSON文件
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)
            
            logger.success(f"JSON数据保存成功: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"保存JSON数据失败: {e}")
            return False

    def _convert_soup_to_markdown(self, soup: BeautifulSoup) -> str:
        """将BeautifulSoup对象转换为Markdown格式"""
        try:
            markdown_lines = []
            
            for element in soup.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'img', 'strong', 'b', 'em', 'i']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(element.name[1])
                    text = element.get_text(strip=True)
                    if text:
                        markdown_lines.append(f"{'#' * level} {text}\n")
                        
                elif element.name == 'img':
                    src = element.get('src', '')
                    alt = element.get('alt', '图片')
                    if src:
                        markdown_lines.append(f"![{alt}]({src})\n")
                        
                elif element.name in ['strong', 'b']:
                    text = element.get_text(strip=True)
                    if text:
                        markdown_lines.append(f"**{text}**")
                        
                elif element.name in ['em', 'i']:
                    text = element.get_text(strip=True)
                    if text:
                        markdown_lines.append(f"*{text}*")
                        
                elif element.name in ['p', 'div']:
                    text = element.get_text(strip=True)
                    if text and len(text) > 5:  # 过滤太短的内容
                        markdown_lines.append(f"{text}\n")
            
            return '\n'.join(markdown_lines)
            
        except Exception as e:
            logger.error(f"转换Markdown失败: {e}")
            return soup.get_text(strip=True) if soup else ""




def main():
    """测试工具"""
    # 示例URL
    test_url = "https://mp.weixin.qq.com/s/LiS8ytwKsKP9yAHGp96G_Q"
    
    scraper = SimpleUrlScraper(headless=False)  # 使用可见模式便于处理验证码
    
    try:
        # 测试PDF生成
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_path = f"output/pdf/test_article_{timestamp}.pdf"
        
        if scraper.save_as_pdf(test_url, pdf_path):
            logger.success(f"PDF测试成功: {pdf_path}")
        
        # 测试完整HTML保存
        html_path = f"output/html/test_article_{timestamp}.html"
        
        if scraper.save_complete_html(test_url, html_path):
            logger.success(f"完整HTML测试成功: {html_path}")
        
        # 测试基本文章信息提取
        article = scraper.extract_article_info(test_url)
        
        if 'error' not in article:
            logger.success(f"文章信息提取成功: {article.get('title')}")
            
        else:
            logger.error(f"处理失败: {article.get('error')}")
            
    except KeyboardInterrupt:
        logger.info("用户中断操作")
    except Exception as e:
        logger.error(f"程序执行出错: {e}")
    finally:
        scraper.cleanup()


if __name__ == "__main__":
    main() 