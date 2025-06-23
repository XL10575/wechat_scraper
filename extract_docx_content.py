#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取DOCX文档内容工具
提取文字和图片，便于复制到飞书知识库
"""

import os
import sys
from pathlib import Path
from loguru import logger
import zipfile
import xml.etree.ElementTree as ET
from io import BytesIO
import base64

try:
    from docx import Document
    from docx.document import Document as _Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 库未安装，将使用基础XML解析")

def extract_docx_content(docx_path: str):
    """提取DOCX文档的内容
    
    Args:
        docx_path: DOCX文件路径
        
    Returns:
        Dict: 包含文本内容和图片信息的字典
    """
    try:
        logger.info(f"📖 开始提取DOCX内容: {os.path.basename(docx_path)}")
        
        if not os.path.exists(docx_path):
            logger.error(f"❌ 文件不存在: {docx_path}")
            return None
        
        if DOCX_AVAILABLE:
            return extract_with_python_docx(docx_path)
        else:
            return extract_with_xml_parser(docx_path)
            
    except Exception as e:
        logger.error(f"提取DOCX内容失败: {e}")
        return None

def extract_with_python_docx(docx_path: str):
    """使用python-docx库提取内容"""
    try:
        doc = Document(docx_path)
        
        content = {
            'title': os.path.splitext(os.path.basename(docx_path))[0],
            'paragraphs': [],
            'images': [],
            'tables': [],
            'full_text': ''
        }
        
        logger.info("📝 提取段落内容...")
        
        # 提取段落内容
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content['paragraphs'].append({
                    'text': text,
                    'style': para.style.name if para.style else 'Normal'
                })
        
        # 提取表格内容
        logger.info("📊 提取表格内容...")
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            if table_data:
                content['tables'].append({
                    'index': table_idx,
                    'data': table_data
                })
        
        # 提取图片（需要解析关系文件）
        logger.info("🖼️ 提取图片信息...")
        images = extract_images_from_docx(docx_path)
        content['images'] = images
        
        # 合并所有文本
        all_text = []
        for para in content['paragraphs']:
            all_text.append(para['text'])
        
        for table in content['tables']:
            for row in table['data']:
                all_text.extend(row)
        
        content['full_text'] = '\n\n'.join(filter(None, all_text))
        
        logger.success(f"✅ 内容提取完成:")
        logger.info(f"📄 段落数: {len(content['paragraphs'])}")
        logger.info(f"📊 表格数: {len(content['tables'])}")
        logger.info(f"🖼️ 图片数: {len(content['images'])}")
        logger.info(f"📝 总字符数: {len(content['full_text'])}")
        
        return content
        
    except Exception as e:
        logger.error(f"python-docx提取失败: {e}")
        return None

def extract_images_from_docx(docx_path: str):
    """从DOCX文件中提取图片"""
    images = []
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_file:
            # 查找图片文件
            image_files = [f for f in zip_file.namelist() 
                          if f.startswith('word/media/') and 
                          any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'])]
            
            logger.info(f"🔍 发现 {len(image_files)} 个图片文件")
            
            for img_file in image_files:
                try:
                    # 读取图片数据
                    img_data = zip_file.read(img_file)
                    
                    # 获取文件扩展名
                    _, ext = os.path.splitext(img_file)
                    
                    # 保存图片到临时目录
                    temp_dir = "temp_images"
                    os.makedirs(temp_dir, exist_ok=True)
                    
                    img_filename = f"image_{len(images)+1}{ext}"
                    img_path = os.path.join(temp_dir, img_filename)
                    
                    with open(img_path, 'wb') as f:
                        f.write(img_data)
                    
                    images.append({
                        'filename': img_filename,
                        'path': img_path,
                        'size': len(img_data),
                        'format': ext[1:].upper()
                    })
                    
                    logger.info(f"💾 保存图片: {img_filename} ({len(img_data)} bytes)")
                    
                except Exception as e:
                    logger.warning(f"处理图片失败 {img_file}: {e}")
                    continue
                    
    except Exception as e:
        logger.error(f"提取图片失败: {e}")
        
    return images

def extract_with_xml_parser(docx_path: str):
    """使用基础XML解析器提取内容"""
    try:
        logger.info("使用基础XML解析器...")
        
        content = {
            'title': os.path.splitext(os.path.basename(docx_path))[0],
            'paragraphs': [],
            'images': [],
            'tables': [],
            'full_text': ''
        }
        
        with zipfile.ZipFile(docx_path, 'r') as zip_file:
            # 读取主文档
            try:
                doc_xml = zip_file.read('word/document.xml')
                root = ET.fromstring(doc_xml)
                
                # 定义命名空间
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # 提取段落
                paragraphs = root.findall('.//w:p', namespaces)
                
                for para in paragraphs:
                    text_elements = para.findall('.//w:t', namespaces)
                    para_text = ''.join(elem.text or '' for elem in text_elements)
                    
                    if para_text.strip():
                        content['paragraphs'].append({
                            'text': para_text.strip(),
                            'style': 'Normal'
                        })
                
                # 合并文本
                content['full_text'] = '\n\n'.join(para['text'] for para in content['paragraphs'])
                
                # 提取图片
                content['images'] = extract_images_from_docx(docx_path)
                
                logger.success(f"✅ XML解析完成:")
                logger.info(f"📄 段落数: {len(content['paragraphs'])}")
                logger.info(f"🖼️ 图片数: {len(content['images'])}")
                logger.info(f"📝 总字符数: {len(content['full_text'])}")
                
                return content
                
            except Exception as e:
                logger.error(f"XML解析失败: {e}")
                return None
                
    except Exception as e:
        logger.error(f"基础XML解析失败: {e}")
        return None

def save_content_for_copy(content, output_file="docx_content_for_copy.md"):
    """将提取的内容保存为Markdown格式，便于复制"""
    try:
        logger.info(f"💾 保存内容到: {output_file}")
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"# {content['title']}\n\n")
            f.write("---\n\n")
            
            # 写入段落内容
            if content['paragraphs']:
                f.write("## 文档内容\n\n")
                for para in content['paragraphs']:
                    f.write(f"{para['text']}\n\n")
            
            # 写入表格
            if content['tables']:
                f.write("## 表格内容\n\n")
                for table_idx, table in enumerate(content['tables']):
                    f.write(f"### 表格 {table_idx + 1}\n\n")
                    
                    if table['data']:
                        # 写入表格头
                        headers = table['data'][0]
                        f.write("| " + " | ".join(headers) + " |\n")
                        f.write("| " + " | ".join(["---"] * len(headers)) + " |\n")
                        
                        # 写入表格数据
                        for row in table['data'][1:]:
                            f.write("| " + " | ".join(row) + " |\n")
                        f.write("\n")
            
            # 写入图片信息
            if content['images']:
                f.write("## 图片信息\n\n")
                f.write("以下图片已保存到 temp_images 文件夹，请手动复制到飞书中：\n\n")
                
                for img in content['images']:
                    f.write(f"- **{img['filename']}** ({img['format']}, {img['size']} bytes)\n")
                    f.write(f"  - 路径: `{img['path']}`\n\n")
        
        logger.success(f"✅ 内容已保存到: {output_file}")
        logger.info("📋 您可以打开此文件复制内容到飞书知识库中")
        
        return output_file
        
    except Exception as e:
        logger.error(f"保存内容失败: {e}")
        return None

def main():
    """主函数"""
    logger.info("=" * 60)
    logger.info("DOCX内容提取工具")
    logger.info("=" * 60)
    
    # 查找DOCX文件
    docx_path = "output/auto_download/更新公告.docx"
    
    if not os.path.exists(docx_path):
        logger.error(f"❌ 找不到文件: {docx_path}")
        
        # 列出可用文件
        auto_download_dir = "output/auto_download"
        if os.path.exists(auto_download_dir):
            files = [f for f in os.listdir(auto_download_dir) if f.endswith('.docx')]
            if files:
                logger.info("📁 可用的DOCX文件:")
                for f in files:
                    logger.info(f"  - {f}")
                
                # 使用第一个找到的DOCX文件
                docx_path = os.path.join(auto_download_dir, files[0])
                logger.info(f"📖 使用文件: {docx_path}")
            else:
                logger.error("❌ 未找到任何DOCX文件")
                return
        else:
            logger.error("❌ auto_download目录不存在")
            return
    
    # 提取内容
    content = extract_docx_content(docx_path)
    
    if content:
        # 保存为便于复制的格式
        output_file = save_content_for_copy(content)
        
        if output_file:
            logger.info("\n" + "=" * 60)
            logger.success("🎉 内容提取完成！")
            logger.info("📋 复制步骤:")
            logger.info("1. 打开生成的Markdown文件查看内容")
            logger.info(f"2. 复制文本内容到飞书知识库: https://thedream.feishu.cn/wiki/R8mSwJ48piDjMwkk0I0cfISNnqb")
            logger.info("3. 手动上传 temp_images 文件夹中的图片到飞书")
            logger.info("=" * 60)
    else:
        logger.error("❌ 内容提取失败")

if __name__ == "__main__":
    main() 