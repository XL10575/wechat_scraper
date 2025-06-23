#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建测试用的DOCX文件
"""

import os
from docx import Document
from docx.shared import Pt
from loguru import logger


def create_test_docx(output_path=None):
    """创建一个简单的测试DOCX文件"""
    try:
        # 创建输出目录
        output_dir = "output/test_docx"
        os.makedirs(output_dir, exist_ok=True)
        
        # 创建文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('测试文档 - DOCX导入为飞书云文档', 0)
        
        # 添加段落
        doc.add_paragraph('这是一个测试文档，用于验证DOCX导入为飞书云文档的功能。')
        
        # 添加二级标题
        doc.add_heading('功能测试', level=1)
        
        # 添加更多内容
        doc.add_paragraph('本文档包含以下内容：')
        
        # 添加列表
        p = doc.add_paragraph('1. 标题测试')
        p = doc.add_paragraph('2. 段落测试')
        p = doc.add_paragraph('3. 格式测试')
        
        # 添加另一个标题
        doc.add_heading('导入流程', level=1)
        
        # 添加步骤说明
        doc.add_paragraph('新的三步导入流程：')
        doc.add_paragraph('步骤一：上传素材文件到飞书云空间')
        doc.add_paragraph('步骤二：创建导入任务（type=docs）')
        doc.add_paragraph('步骤三：查询导入结果并获取飞书云文档URL')
        
        # 保存文件
        if output_path:
            file_path = output_path
        else:
            file_path = os.path.join(output_dir, "测试DOCX导入功能.docx")
        doc.save(file_path)
        
        logger.success(f"✅ 测试DOCX文件创建成功: {file_path}")
        return file_path
        
    except Exception as e:
        logger.error(f"创建测试DOCX文件失败: {e}")
        return None


def main():
    """主函数"""
    logger.info("🚀 开始创建测试DOCX文件")
    
    file_path = create_test_docx()
    
    if file_path:
        logger.success("🎉 测试文件创建完成")
        logger.info(f"📁 文件位置: {file_path}")
        logger.info("💡 现在可以用这个文件测试导入功能")
    else:
        logger.error("❌ 测试文件创建失败")


if __name__ == "__main__":
    main() 